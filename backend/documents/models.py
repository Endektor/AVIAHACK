from django.db import models
from django.db.models.signals import post_save, pre_save
from django.dispatch import receiver
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.core.files import File
from django.core.files.base import ContentFile

import io
import os
import docx
import qrcode
from itertools import product
from docx2pdf import convert
import datetime


class Pattern(models.Model):
    name = models.CharField(max_length=128)
    file = models.FileField(upload_to='patterns/')

    def __str__(self):
        return self.name


class Field(models.Model):
    key = models.CharField(max_length=128)
    value = models.TextField()
    document = models.ForeignKey('Document', on_delete=models.CASCADE, related_name='document_field')

    def __str__(self):
        return self.key


def get_name():
    return datetime.datetime.today().strftime("%Y-%m-%d_%H-%M-%S")


class Document(models.Model):
    name = models.CharField(max_length=128, blank=True, default=get_name)
    pattern = models.ForeignKey('Pattern', on_delete=models.PROTECT, related_name='pattern_document')
    document = models.FileField(upload_to='documents/', blank=True, null=True)
    qr = models.ImageField(upload_to='qr/', blank=True, null=True)

    def __str__(self):
        return self.name

    def save(self, *args, **kwargs):
        data = 'id:' + str(self.id)
        img = qrcode.make(data)
        file_bytes = io.BytesIO()
        img.save(file_bytes)
        file_bytes.seek(0)
        try:
            os.remove(self.qr.path)
        except:
            pass
        self.qr.save(self.name + '.png', ContentFile(file_bytes.read()), save=False)
        super(Document, self).save(*args, **kwargs)


@receiver(post_save, sender=Field)
def document_creation(sender, *args, **kwargs):
    field_obj = kwargs.get('instance')
    document_obj = field_obj.document
    document = docx.Document(document_obj.pattern.file)
    fields = Field.objects.filter(document=document_obj)
    for field, paragraph in product(fields, document.paragraphs):
        paragraph.text = paragraph.text.replace('<<' + field.key + '>>', field.value)

    file_bytes = io.BytesIO()
    document.save(file_bytes)
    file_bytes.seek(0)
    try:
        os.remove(document_obj.document.path)
    except:
        pass
    document_obj.document.save(document_obj.name + '.docx', ContentFile(file_bytes.read()), save=True)
    convert(document_obj.document.path)
    with open(document_obj.document.path[:-4] + 'pdf', 'rb') as f:
        document_obj.document.save(document_obj.name + '.pdf', File(f), save=True)
